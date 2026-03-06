#!/usr/bin/env python3
"""
Add Extended Captions Section to DOCX

This script adds a SECTION 05 "Extended Captions" table to the end of a 
Content Schedule DOCX file. The table provides a place to store long captions
that would clutter the main posting schedule table.

Usage:
    python add_extended_captions.py input.docx                     # overwrites input
    python add_extended_captions.py input.docx output.docx         # creates new file

The script is idempotent - it won't duplicate the section if it already exists.
"""

import sys
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def check_section_exists(doc):
    """
    Check if Extended Captions section already exists in document.
    Returns True if found.
    """
    for para in doc.paragraphs:
        text_lower = para.text.lower().strip()
        if "extended caption" in text_lower or "section 05" in text_lower:
            return True
    
    # Also check tables for the characteristic header
    for table in doc.tables:
        if table.rows:
            first_row = table.rows[0]
            if len(first_row.cells) >= 2:
                headers = [cell.text.lower().strip() for cell in first_row.cells[:2]]
                if "post title" in headers[0] and "caption" in headers[1]:
                    return True
    
    return False


def set_cell_shading(cell, hex_color):
    """Apply background shading to a table cell."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), hex_color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_table_borders(table):
    """Set borders on a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    
    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def add_extended_captions_section(doc):
    """
    Add SECTION 05 Extended Captions with an empty template table.
    """
    # Add some spacing before the new section
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Section number
    section_num = doc.add_paragraph()
    section_num.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = section_num.add_run("SECTION 05")
    run.bold = True
    run.font.size = Pt(14)
    
    # Section title
    section_title = doc.add_paragraph()
    section_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = section_title.add_run("Extended Captions")
    run.bold = True
    run.font.size = Pt(18)
    
    # Description
    description = doc.add_paragraph()
    description.add_run(
        "Optional: For longer captions that don't fit well in the posting schedule table. "
        "The Post Title must match exactly with SECTION 01."
    )
    
    # Add spacing
    doc.add_paragraph()
    
    # Create the 2-column table
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set table borders manually (since Table Grid style may not exist)
    set_table_borders(table)
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = "Post Title"
    header_cells[1].text = "Extended Caption"
    
    # Style header row
    for cell in header_cells:
        # Make text bold
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
        # Light gray background
        set_cell_shading(cell, "E8E8E8")
    
    # Template row (placeholder that user will replace)
    data_cells = table.rows[1].cells
    data_cells[0].text = "[Exact Post Title from SECTION 01]"
    data_cells[1].text = "[Paste your full caption here - can be multiple paragraphs]"
    
    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(2.5)
        row.cells[1].width = Inches(4.5)
    
    # Add usage note
    doc.add_paragraph()
    note = doc.add_paragraph()
    note_run = note.add_run("Tip: ")
    note_run.italic = True
    note.add_run(
        "To add more rows, click in the last cell and press Tab. "
        "Leave the Caption column in SECTION 01 empty or use [EXT] marker."
    )
    
    return True


def main():
    if len(sys.argv) < 2:
        print("Usage: python add_extended_captions.py input.docx [output.docx]")
        print("\nAdds an Extended Captions section to your content schedule DOCX.")
        print("If no output file is specified, creates input_with_captions.docx")
        sys.exit(1)
    
    input_file = sys.argv[1]
    
    # Determine output file
    if len(sys.argv) >= 3:
        output_file = sys.argv[2]
    else:
        base, ext = os.path.splitext(input_file)
        output_file = f"{base}_with_captions{ext}"
    
    # Validate input exists
    if not os.path.exists(input_file):
        print(f"❌ Error: File '{input_file}' not found.")
        sys.exit(1)
    
    print(f"📄 Reading: {input_file}")
    
    # Load document
    doc = Document(input_file)
    
    # Check if section already exists
    if check_section_exists(doc):
        print("⚠️  Extended Captions section already exists in this document.")
        print("    No changes made.")
        sys.exit(0)
    
    # Add the section
    print("📝 Adding Extended Captions section...")
    add_extended_captions_section(doc)
    
    # Save
    doc.save(output_file)
    
    print(f"\n✅ Saved: {output_file}")
    print("\nNext steps:")
    print("  1. Open the DOCX file")
    print("  2. Scroll to SECTION 05 Extended Captions at the end")
    print("  3. Add rows for any posts that need long captions")
    print("  4. Run generate_report.py to create the HTML")


if __name__ == "__main__":
    main()
