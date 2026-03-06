"""
Generate UGC Management SOP (DOCX) and UGC Tracking Spreadsheet (XLSX)
from the Smoothie Bar UGC SOP content.
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import openpyxl
from openpyxl.styles import (
    PatternFill, Font, Alignment, Border, Side, GradientFill
)
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.datavalidation import DataValidation
from openpyxl.formatting.rule import Rule
from openpyxl.styles.differential import DifferentialStyle

OUTPUT_DIR = r"c:\dev\projects\Smoothie_Bar_IG\SOP & Rules"

# ─────────────────────────────────────────────────────────────────────────────
# COLOR PALETTE
# ─────────────────────────────────────────────────────────────────────────────
DARK       = "1A1A1A"
MID_GRAY   = "6B6B6B"
LIGHT_BG   = "FAF9F7"
WARM       = "C4A484"
BORDER_CLR = "E8E6E3"
WHITE      = "FFFFFF"

# Status colors (hex, no #)
STATUS_COLORS = {
    "CONTACTED":          ("FFF9C4", "5D4E00"),   # soft yellow
    "CONFIRMED":          ("BBDEFB", "0D3C61"),   # light blue
    "SHIPPED":            ("FFE0B2", "6D3200"),   # light orange
    "DELIVERED":          ("E8D5F5", "4A1480"),   # light purple
    "CONTENT RECEIVED":   ("C8E6C9", "1B5E20"),   # mint green
    "RESHOOT REQUESTED":  ("FFD180", "5D3A00"),   # amber
    "POSTED":             ("A5D6A7", "1B5E20"),   # green
    "NO RESPONSE":        ("FFCDD2", "7F0000"),   # pink/red
    "DECLINED":           ("E0E0E0", "424242"),   # light gray
}

TIER_COLORS = {
    "COLD":     ("CFD8DC", "263238"),   # cool gray-blue
    "WARM":     ("FFE0B2", "4E2600"),   # light orange
    "RELIABLE": ("C8E6C9", "1B5E20"),   # soft green
    "VIP":      ("FFF176", "3E2800"),   # gold
}


# ─────────────────────────────────────────────────────────────────────────────
# DOCX HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def hex_to_rgb(hex_str):
    h = hex_str.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def set_paragraph_border_bottom(paragraph, color="E8E6E3", size=4):
    """Add a bottom border to a paragraph via XML."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_shaded_paragraph(doc, text, shade_hex="F3F1EE", left_indent=0.3):
    """Add a paragraph with a background shade (for DM template boxes)."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), shade_hex)
    pPr.append(shd)
    p.paragraph_format.left_indent = Inches(left_indent)
    p.paragraph_format.right_indent = Inches(left_indent)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = hex_to_rgb(DARK)
    return p


def add_section_header(doc, text, level=1):
    """Add a styled section header."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    if level == 1:
        # Full-width dark bar
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1A1A1A")
        pPr.append(shd)
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.left_indent = Inches(0.2)
        run = p.add_run(text.upper())
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = hex_to_rgb(WHITE)
        run.font.name = "Calibri"
    elif level == 2:
        # Warm accent underline style
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        run.font.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = hex_to_rgb("8B6E4E")
        run.font.name = "Calibri"
        set_paragraph_border_bottom(p, color="C4A484", size=6)
    return p


def add_step(doc, number, title, body):
    """Add a numbered step with title and body."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0)

    # Number badge via bold colored run
    num_run = p.add_run(f"  {number:02d}  ")
    num_run.font.bold = True
    num_run.font.size = Pt(10)
    num_run.font.color.rgb = hex_to_rgb(WHITE)
    num_run.font.name = "Calibri"

    # Shade the number via XML on its run
    rPr = num_run._r.get_or_add_rPr()
    shd2 = OxmlElement("w:shd")
    shd2.set(qn("w:val"), "clear")
    shd2.set(qn("w:color"), "auto")
    shd2.set(qn("w:fill"), "C4A484")
    rPr.append(shd2)

    title_run = p.add_run(f"  {title}")
    title_run.font.bold = True
    title_run.font.size = Pt(10.5)
    title_run.font.color.rgb = hex_to_rgb(DARK)
    title_run.font.name = "Calibri"

    body_p = doc.add_paragraph(body)
    body_p.paragraph_format.left_indent = Inches(0.45)
    body_p.paragraph_format.space_before = Pt(2)
    body_p.paragraph_format.space_after = Pt(6)
    for run in body_p.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = hex_to_rgb(MID_GRAY)
        run.font.name = "Calibri"


def add_checklist_item(doc, title, body=""):
    """Add a daily checklist item."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)

    box_run = p.add_run("  \u25a1  ")
    box_run.font.bold = True
    box_run.font.size = Pt(12)
    box_run.font.color.rgb = hex_to_rgb(WARM)

    title_run = p.add_run(title)
    title_run.font.bold = True
    title_run.font.size = Pt(10.5)
    title_run.font.color.rgb = hex_to_rgb(DARK)
    title_run.font.name = "Calibri"

    if body:
        body_p = doc.add_paragraph(body)
        body_p.paragraph_format.left_indent = Inches(0.7)
        body_p.paragraph_format.space_before = Pt(1)
        body_p.paragraph_format.space_after = Pt(4)
        for run in body_p.runs:
            run.font.size = Pt(9.5)
            run.font.color.rgb = hex_to_rgb(MID_GRAY)
            run.font.name = "Calibri"


def add_dm_template(doc, label, template_text):
    """Add a labeled DM template in a shaded box."""
    # Label
    lp = doc.add_paragraph()
    lp.paragraph_format.space_before = Pt(10)
    lp.paragraph_format.space_after = Pt(0)
    lr = lp.add_run(label.upper())
    lr.font.bold = True
    lr.font.size = Pt(9)
    lr.font.color.rgb = hex_to_rgb("8B6E4E")
    lr.font.name = "Calibri"

    # Template box
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F5F3F0")
    pPr.append(shd)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(template_text)
    run.font.size = Pt(9.5)
    run.font.color.rgb = hex_to_rgb(DARK)
    run.font.name = "Calibri"
    run.font.italic = True


def add_bullet(doc, text, bold_prefix=None):
    """Add a bullet point, optionally with a bold prefix."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.first_line_indent = Inches(-0.2)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

    dot_run = p.add_run("  \u2022  ")
    dot_run.font.color.rgb = hex_to_rgb(WARM)
    dot_run.font.size = Pt(10)

    if bold_prefix:
        bp = p.add_run(bold_prefix + " ")
        bp.font.bold = True
        bp.font.size = Pt(10)
        bp.font.color.rgb = hex_to_rgb(DARK)
        bp.font.name = "Calibri"

    body_run = p.add_run(text)
    body_run.font.size = Pt(10)
    body_run.font.color.rgb = hex_to_rgb(MID_GRAY)
    body_run.font.name = "Calibri"


def add_ref_table(doc, columns):
    """Add a clean reference table for spreadsheet columns."""
    # Split into rows of 4
    chunk = 4
    chunked = [columns[i:i+chunk] for i in range(0, len(columns), chunk)]
    table = doc.add_table(rows=1 + len(chunked), cols=chunk)
    table.style = "Table Grid"

    # Header row
    for j, hdr in enumerate(["Column", "Column", "Column", "Column"]):
        pass  # We'll just fill data

    # Fill data
    for i, row_data in enumerate(chunked):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_text
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Calibri"
                    run.font.color.rgb = hex_to_rgb(DARK)
            # Shade alternate rows
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "F5F3F0" if i % 2 == 0 else "FEFEFE")
            tcPr.append(shd)

        # Fill empty cells in last row if needed
        if len(row_data) < chunk:
            for j in range(len(row_data), chunk):
                row.cells[j].text = ""

    return table


# ─────────────────────────────────────────────────────────────────────────────
# BUILD DOCX
# ─────────────────────────────────────────────────────────────────────────────

def build_docx():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # ── TITLE BLOCK ──────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)

    t_run = title_p.add_run("UGC MANAGEMENT SOP")
    t_run.font.bold = True
    t_run.font.size = Pt(22)
    t_run.font.color.rgb = hex_to_rgb(DARK)
    t_run.font.name = "Calibri"

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_before = Pt(0)
    sub_p.paragraph_format.space_after = Pt(2)
    sub_run = sub_p.add_run("Smoothie Bar Infusions  |  User-Generated Content Program")
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = hex_to_rgb(WARM)
    sub_run.font.name = "Calibri"

    date_p = doc.add_paragraph()
    date_p.paragraph_format.space_before = Pt(0)
    date_p.paragraph_format.space_after = Pt(16)
    set_paragraph_border_bottom(date_p, color="C4A484", size=8)
    date_run = date_p.add_run("Internal Use Only  —  Updated March 2026")
    date_run.font.size = Pt(9)
    date_run.font.color.rgb = hex_to_rgb(MID_GRAY)
    date_run.font.name = "Calibri"

    # ── SECTION 1: 11-STEP PROCESS ───────────────────────────────────────────
    add_section_header(doc, "11-Step UGC Process", level=1)

    steps = [
        ("FIND & VET",
         "Scroll accounts like @wholemelts and similar brands. Look for creators already posting UGC-style cannabis content. Before reaching out, confirm: they look 21+, their content style fits our brand, and their audience looks real (not botted)."),
        ("DM FROM ALT",
         "Follow from alt account. Send the initial outreach message (see DM Templates) and attach the UGC Guidelines doc with the DM."),
        ("GET CONFIRMATION",
         "Wait for a yes. Get their shipping address. Do NOT send product to anyone who hasn't confirmed and acknowledged the guidelines."),
        ("SUBMIT ORDER & LOG",
         "Send the order to merch chat for fulfillment. Log on the UGC spreadsheet: creator handle, date contacted, date confirmed, what's being sent, shipping address. Also mark this is a UGC account."),
        ("SEND TRACKING",
         "Once shipped, DM them the tracking number on IG."),
        ("CONFIRM DELIVERY",
         "Follow up to confirm the bag arrived. Remind them of the content deadline (7-14 days from delivery)."),
        ("FOLLOW UP ON CONTENT",
         "Day 7: gentle nudge. Day 14: firmer check-in. Day 21: if nothing, mark as 'No Response' and move on."),
        ("SCREEN CONTENT",
         "Review against the UGC Guidelines. Check for: no consumption shown, no minors, no health claims, no open product, no flaggable language/music."),
        ("REPOST OR REQUEST RESHOOT",
         "Passes guidelines: Confirm creator is OK with repost, then repost on main account using IG's native repost/collab tools. Breaks guidelines: Ask if they can reshoot to fit the rules. Can't reshoot: Decide case-by-case whether to download and repost with edits, or skip it. Contact internal team group chat for opinions if necessary."),
        ("UPDATE SPREADSHEET",
         "Mark each creator's status: COMPLETED, NO POST, NO RESPONSE, RESHOOT REQUESTED, IN PROGRESS. Keep this current — check it every day."),
        ("TRACK PERFORMANCE",
         "Check metrics on every reposted UGC piece: views, likes, saves, shares, profile visits. Flag top performers. Tier creators in the spreadsheet (Cold / Warm / Reliable / VIP). VIPs get priority on future drops and bigger incentives."),
    ]

    for i, (title, body) in enumerate(steps, 1):
        add_step(doc, i, title, body)

    # ── SECTION 2: DAILY CHECKLIST ───────────────────────────────────────────
    doc.add_paragraph()
    add_section_header(doc, "Daily Checklist", level=1)

    intro_p = doc.add_paragraph("Complete these every day, in this order.")
    intro_p.paragraph_format.space_before = Pt(6)
    intro_p.paragraph_format.space_after = Pt(8)
    for run in intro_p.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = hex_to_rgb(MID_GRAY)
        run.font.name = "Calibri"
        run.font.italic = True

    checklist = [
        ("CHECK SPREADSHEET",
         "Open the UGC spreadsheet. Review all active creators and their current status. This is your dashboard — everything flows from here."),
        ("CHECK DMs",
         "Go through the alt account inbox. Look for: new replies, shipping address confirmations, content submissions, questions about guidelines."),
        ("PROCESS NEW CONFIRMATIONS",
         "Any creator who said yes and sent their address: submit their order to merch chat, log the details on the spreadsheet (handle, date confirmed, items sent, address)."),
        ("SEND TRACKING",
         "Check for any new tracking numbers from merch. DM tracking to the relevant creators. Update spreadsheet."),
        ("FOLLOW UP ON DELIVERIES",
         "Check tracking on all shipped packages. If delivered, DM creator to confirm receipt and remind them of the content deadline. Update spreadsheet."),
        ("FOLLOW UP ON PENDING CONTENT",
         "Nudge creators who are past due: day 7 = gentle reminder, day 14 = direct ask, day 21 = mark as No Response and move on. Update spreadsheet."),
        ("SCREEN INCOMING CONTENT",
         "Review new content submissions against the UGC Guidelines. Pass: get repost permission and repost via IG native tools. Fail: ask for reshoot or make a call on whether to skip. Update spreadsheet."),
        ("CHECK METRICS",
         "Pull numbers on all live UGC reposts: views, likes, saves, shares. Note standout performers. Update creator tiers if anyone's earned a bump."),
        ("SCOUT NEW CREATORS",
         "Spend 15-20 min scrolling relevant accounts. Vet promising creators (21+, good content style, real audience). DM from alt with pitch + guidelines. Log on spreadsheet as 'CONTACTED.'"),
        ("END OF DAY: SPREADSHEET AUDIT",
         "Make sure every active creator has an up-to-date status. No row should be stale."),
    ]

    for title, body in checklist:
        add_checklist_item(doc, title, body)

    # ── SECTION 3: CREATOR GUIDELINES ────────────────────────────────────────
    doc.add_paragraph()
    add_section_header(doc, "Guidelines for Content Creators", level=1)

    add_section_header(doc, "What We Want", level=2)
    intro2 = doc.add_paragraph(
        "Your honest reaction. Unboxing, first impressions, aesthetic shots, packaging closeups, lifestyle vibes. "
        "Be yourself — we picked you because we like your style."
    )
    intro2.paragraph_format.space_before = Pt(6)
    intro2.paragraph_format.space_after = Pt(8)
    for run in intro2.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = hex_to_rgb(MID_GRAY)
        run.font.name = "Calibri"

    add_section_header(doc, "Timeline & Reposting", level=2)
    timeline_p = doc.add_paragraph(
        "We'd love to have your content within 2 weeks of receiving the bag. If you need more time, just let us know. "
        "By sending us your content, you're giving us permission to repost on our pages with credit to you. "
        "We'll always tag you and use IG's collab tools when possible."
    )
    timeline_p.paragraph_format.space_before = Pt(6)
    timeline_p.paragraph_format.space_after = Pt(8)
    for run in timeline_p.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = hex_to_rgb(MID_GRAY)
        run.font.name = "Calibri"

    add_section_header(doc, "Do Not Show or Include", level=2)
    dont_list = [
        "No consumption of any kind (smoking, vaping, eating edibles, dabbing) — not even implied",
        "No minors, and no one who could reasonably be mistaken for under 21",
        "No health or medical claims (\"cures anxiety,\" \"helps with pain,\" etc.)",
        "No driving, operating machinery, or any activity where impairment would be dangerous",
        "No prices, discounts, or promotional offers",
        "No sales or implied sales",
    ]
    for item in dont_list:
        add_bullet(doc, item)

    add_section_header(doc, "Always Do", level=2)
    do_list = [
        "All talent must be visibly 21+ and be able to verify age if asked",
        "Keep it clean and professional",
        "Focus on lifestyle, branding, aesthetics, and the unboxing/packaging experience",
        "Include disclaimers or age-gate language for content",
        "Tag us on content (@smoothiebarinfusions) and use \"Smoothie Bar: Blend 3.0\" in video titles",
        "Keep everything platform-safe — assume the content needs to survive on Instagram and TikTok",
    ]
    for item in do_list:
        add_bullet(doc, item)

    add_section_header(doc, "Gray Areas to Avoid", level=2)
    gray_list = [
        "Don't show or reference alcohol or other substances alongside the product",
        "Avoid cannabis terms or slang that platforms auto-flag, especially in caption text",
        "Don't use music you don't have a license for",
    ]
    for item in gray_list:
        add_bullet(doc, item)

    # ── SECTION 4: DM TEMPLATES ───────────────────────────────────────────────
    doc.add_paragraph()
    add_section_header(doc, "DM Templates", level=1)

    templates = [
        ("Initial Outreach",
         "Hey [name]! We came across your page and love your content — the aesthetic is exactly what we're going for. "
         "We're [Brand] and we'd love to send you a goodie bag in exchange for a product review or unboxing video, just your honest take. "
         "Interested? I'll send over our content guidelines so you know what works for reposting."),
        ("After They Say Yes",
         "Awesome! Here are our content guidelines [attach]. Give those a quick read when you get a chance — the main thing is "
         "no consumption on camera and keeping everything clean and professional as possible. "
         "Drop your shipping address and we'll get your bag out ASAP."),
        ("Tracking",
         "Hey! Your bag just shipped — here's your tracking: [#]. Should be there by [date]. "
         "Once you get it, just shoot me a message and take your time with the content. We just ask for something within 2 weeks of delivery."),
        ("Day 7 Nudge",
         "Hey [name]! Just checking in — did you get a chance to check out the bag? No rush, just wanted to make sure everything arrived good. "
         "Looking forward to seeing what you put together!"),
        ("Day 14 Follow-Up",
         "Hey! Circling back on this — any update on the content? We'd love to get your video up soon. "
         "Let me know if you have any questions about the guidelines or need anything from us."),
        ("Day 21 Final",
         "Hey [name], just wanted to follow up one last time. Totally understand if the timing didn't work out — "
         "if you're still down to post, we're still interested. If not, no worries at all."),
        ("Repost Permission",
         "This came out great! We'd love to repost this on our main page — cool with you? "
         "We'll tag you and use IG's collab feature so it shows on both our pages."),
        ("Reshoot Request",
         "Hey, love the effort on this! One thing — [specific issue, e.g. 'there's a quick shot of the product being used that we can't have in the final cut']. "
         "Would you be able to do a quick reshoot with that part swapped out? Everything else is perfect. "
         "Happy to walk you through what would work."),
    ]

    for label, text in templates:
        add_dm_template(doc, label, text)

    # ── SECTION 5: SPREADSHEET REFERENCE ─────────────────────────────────────
    doc.add_paragraph()
    add_section_header(doc, "Spreadsheet Reference", level=1)

    ref_p = doc.add_paragraph("Spreadsheet columns (in order):")
    ref_p.paragraph_format.space_before = Pt(6)
    ref_p.paragraph_format.space_after = Pt(6)
    for run in ref_p.runs:
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = hex_to_rgb(DARK)
        run.font.name = "Calibri"

    columns = [
        "Handle", "Name", "Followers", "Date Contacted",
        "Date Confirmed", "Address", "Items Sent", "Date Shipped",
        "Tracking #", "Date Delivered", "Content Deadline", "Date Content Received",
        "Guidelines Pass (Y/N)", "Date Posted", "Post Link", "Views",
        "Likes", "Saves", "Shares", "Status",
        "Creator Tier", "Notes",
    ]
    add_ref_table(doc, columns)

    doc.add_paragraph()
    add_section_header(doc, "Statuses", level=2)
    statuses = ["CONTACTED", "CONFIRMED", "SHIPPED", "DELIVERED",
                "CONTENT RECEIVED", "RESHOOT REQUESTED", "POSTED",
                "NO RESPONSE", "DECLINED"]
    status_descs = {
        "CONTACTED": "Initial DM sent, awaiting response",
        "CONFIRMED": "Creator agreed and gave shipping address",
        "SHIPPED": "Merch order placed and shipped",
        "DELIVERED": "Package confirmed delivered",
        "CONTENT RECEIVED": "Creator submitted their content",
        "RESHOOT REQUESTED": "Content flagged; reshoot requested",
        "POSTED": "Content reviewed and reposted to main",
        "NO RESPONSE": "No reply by day 21 follow-up",
        "DECLINED": "Creator declined or opted out",
    }
    for s in statuses:
        add_bullet(doc, status_descs[s], bold_prefix=s)

    doc.add_paragraph()
    add_section_header(doc, "Creator Tiers", level=2)
    tier_descs = {
        "COLD": "New creator, unproven — first engagement",
        "WARM": "Delivered content, decent quality, still building trust",
        "RELIABLE": "Consistent quality, meets deadlines, brand-aligned",
        "VIP": "Top performer — priority drops, larger incentives, long-term partner",
    }
    for t, desc in tier_descs.items():
        add_bullet(doc, desc, bold_prefix=t)

    # Save
    out_path = os.path.join(OUTPUT_DIR, "UGC Management SOP.docx")
    doc.save(out_path)
    print(f"DOCX saved: {out_path}")
    return out_path


# ─────────────────────────────────────────────────────────────────────────────
# BUILD XLSX
# ─────────────────────────────────────────────────────────────────────────────

def hex_fill(hex_str):
    return PatternFill(start_color=hex_str, end_color=hex_str, fill_type="solid")


def thin_border():
    s = Side(style="thin", color="E8E6E3")
    return Border(left=s, right=s, top=s, bottom=s)


def build_xlsx():
    wb = openpyxl.Workbook()

    # ── SHEET 1: TRACKER ──────────────────────────────────────────────────────
    ws = wb.active
    ws.title = "UGC Tracker"
    ws.sheet_view.showGridLines = False

    columns = [
        ("Handle",                 18),
        ("Name",                   16),
        ("Followers",              12),
        ("Date Contacted",         15),
        ("Date Confirmed",         15),
        ("Address",                28),
        ("Items Sent",             18),
        ("Date Shipped",           15),
        ("Tracking #",             18),
        ("Date Delivered",         15),
        ("Content Deadline",       16),
        ("Date Content Received",  20),
        ("Guidelines Pass (Y/N)",  20),
        ("Date Posted",            14),
        ("Post Link",              30),
        ("Views",                  10),
        ("Likes",                  10),
        ("Saves",                  10),
        ("Shares",                 10),
        ("Status",                 20),
        ("Creator Tier",           14),
        ("Notes",                  35),
    ]

    # Set column widths
    for col_idx, (col_name, col_width) in enumerate(columns, 1):
        ws.column_dimensions[get_column_letter(col_idx)].width = col_width

    # Header row styling
    header_fill = hex_fill("1A1A1A")
    header_font = Font(name="Calibri", size=10, bold=True, color="FFFFFF")
    header_border = thin_border()

    for col_idx, (col_name, _) in enumerate(columns, 1):
        cell = ws.cell(row=1, column=col_idx, value=col_name)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = header_border

    ws.row_dimensions[1].height = 32

    # Sample data rows
    sample_rows = [
        ["@sample_creator", "Jane Doe", 12400, "2026-03-01", "2026-03-02",
         "123 Main St, Los Angeles CA 90001", "Blend 3.0 Goodie Bag", "2026-03-03",
         "1Z999AA10123456784", "2026-03-05", "2026-03-19", "", "",
         "", "", "", "", "", "", "SHIPPED", "WARM", "Follow up tracking"],
        ["@another_creator", "Marcus Lee", 8900, "2026-03-02", "",
         "", "", "", "", "", "", "", "",
         "", "", "", "", "", "", "CONTACTED", "COLD", "Awaiting reply"],
    ]

    data_font = Font(name="Calibri", size=10, color="1A1A1A")
    alt_fill = hex_fill("FAF9F7")
    white_fill = hex_fill("FFFFFF")

    for row_idx, row_data in enumerate(sample_rows, 2):
        fill = white_fill if row_idx % 2 == 0 else alt_fill
        for col_idx, value in enumerate(row_data, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.fill = fill
            cell.font = data_font
            cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=False)
            cell.border = thin_border()

    # Add empty rows ready to fill (rows 4–53)
    for row_idx in range(4, 54):
        fill = white_fill if row_idx % 2 == 0 else alt_fill
        for col_idx in range(1, len(columns) + 1):
            cell = ws.cell(row=row_idx, column=col_idx, value="")
            cell.fill = fill
            cell.font = data_font
            cell.alignment = Alignment(horizontal="left", vertical="center")
            cell.border = thin_border()
        ws.row_dimensions[row_idx].height = 18

    # Freeze top row
    ws.freeze_panes = "A2"

    # Data validation for Status column (col 20)
    status_values = list(STATUS_COLORS.keys())
    status_formula = '"' + ','.join(status_values) + '"'
    dv_status = DataValidation(
        type="list",
        formula1=status_formula,
        allow_blank=True,
        showDropDown=False,
        showErrorMessage=True,
        errorTitle="Invalid Status",
        error="Please select a valid status from the list."
    )
    ws.add_data_validation(dv_status)
    dv_status.sqref = "T2:T200"

    # Data validation for Creator Tier column (col 21)
    tier_values = list(TIER_COLORS.keys())
    tier_formula = '"' + ','.join(tier_values) + '"'
    dv_tier = DataValidation(
        type="list",
        formula1=tier_formula,
        allow_blank=True,
        showDropDown=False,
        showErrorMessage=True,
        errorTitle="Invalid Tier",
        error="Please select a valid creator tier from the list."
    )
    ws.add_data_validation(dv_tier)
    dv_tier.sqref = "U2:U200"

    # Conditional formatting for Status column (T)
    for status, (bg, fg) in STATUS_COLORS.items():
        fill = PatternFill(start_color=bg, end_color=bg, fill_type="solid")
        font = Font(name="Calibri", size=10, bold=True, color=fg)
        dxf = DifferentialStyle(fill=fill, font=font)
        rule = Rule(type="containsText", operator="containsText",
                    text=status, dxf=dxf)
        rule.formula = [f'NOT(ISERROR(SEARCH("{status}",T1)))']
        ws.conditional_formatting.add("T1:T200", rule)

    # Conditional formatting for Creator Tier column (U)
    for tier, (bg, fg) in TIER_COLORS.items():
        fill = PatternFill(start_color=bg, end_color=bg, fill_type="solid")
        font = Font(name="Calibri", size=10, bold=True, color=fg)
        dxf = DifferentialStyle(fill=fill, font=font)
        rule = Rule(type="containsText", operator="containsText",
                    text=tier, dxf=dxf)
        rule.formula = [f'NOT(ISERROR(SEARCH("{tier}",U1)))']
        ws.conditional_formatting.add("U1:U200", rule)

    # ── SHEET 2: LEGEND ───────────────────────────────────────────────────────
    leg = wb.create_sheet(title="Legend")
    leg.sheet_view.showGridLines = False
    leg.column_dimensions["A"].width = 3
    leg.column_dimensions["B"].width = 22
    leg.column_dimensions["C"].width = 42
    leg.column_dimensions["D"].width = 3
    leg.column_dimensions["E"].width = 16
    leg.column_dimensions["F"].width = 42

    def legend_header(ws, row, col, text):
        cell = ws.cell(row=row, column=col, value=text)
        cell.font = Font(name="Calibri", size=11, bold=True, color="1A1A1A")
        cell.fill = hex_fill("F5F3F0")
        cell.alignment = Alignment(horizontal="left", vertical="center")
        cell.border = Border(
            bottom=Side(style="medium", color="C4A484")
        )

    def legend_row(ws, row, col_start, bg, fg, label, description):
        # Color swatch cell
        swatch = ws.cell(row=row, column=col_start, value="")
        swatch.fill = hex_fill(bg)
        swatch.border = thin_border()

        label_cell = ws.cell(row=row, column=col_start + 1, value=label)
        label_cell.font = Font(name="Calibri", size=10, bold=True, color=fg)
        label_cell.fill = hex_fill(bg)
        label_cell.alignment = Alignment(horizontal="left", vertical="center")
        label_cell.border = thin_border()

        desc_cell = ws.cell(row=row, column=col_start + 2, value=description)
        desc_cell.font = Font(name="Calibri", size=10, color="6B6B6B")
        desc_cell.fill = hex_fill("FAFAFA")
        desc_cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
        desc_cell.border = thin_border()

    # Title
    title_cell = leg.cell(row=1, column=2, value="UGC Tracker — Color Legend")
    title_cell.font = Font(name="Calibri", size=14, bold=True, color="1A1A1A")
    title_cell.alignment = Alignment(horizontal="left", vertical="center")
    leg.row_dimensions[1].height = 28
    leg.row_dimensions[2].height = 8  # spacer

    # Status section
    legend_header(leg, 3, 2, "STATUS COLORS")
    legend_header(leg, 3, 3, "")

    status_descriptions = {
        "CONTACTED":          "Initial DM sent, awaiting creator response",
        "CONFIRMED":          "Creator agreed and provided shipping address",
        "SHIPPED":            "Merch order has been shipped out",
        "DELIVERED":          "Package confirmed delivered to creator",
        "CONTENT RECEIVED":   "Creator submitted their content for review",
        "RESHOOT REQUESTED":  "Content flagged; creator asked to reshoot",
        "POSTED":             "Content approved and reposted to main account",
        "NO RESPONSE":        "No reply received by Day 21 follow-up",
        "DECLINED":           "Creator declined or opted out of program",
    }

    for i, (status, (bg, fg)) in enumerate(STATUS_COLORS.items()):
        r = 4 + i
        legend_row(leg, r, 2, bg, fg, status, status_descriptions[status])
        leg.row_dimensions[r].height = 22

    # Spacer
    leg.row_dimensions[13].height = 12

    # Tier section
    legend_header(leg, 14, 2, "CREATOR TIER COLORS")
    legend_header(leg, 14, 3, "")

    tier_descriptions = {
        "COLD":     "New creator, first engagement — unproven",
        "WARM":     "Delivered content; decent quality, still building trust",
        "RELIABLE": "Consistent quality, meets deadlines, brand-aligned",
        "VIP":      "Top performer — priority drops, larger incentives, long-term partner",
    }

    for i, (tier, (bg, fg)) in enumerate(TIER_COLORS.items()):
        r = 15 + i
        legend_row(leg, r, 2, bg, fg, tier, tier_descriptions[tier])
        leg.row_dimensions[r].height = 22

    # Additional notes section
    leg.row_dimensions[20].height = 12
    notes_hdr = leg.cell(row=21, column=2, value="NOTES")
    notes_hdr.font = Font(name="Calibri", size=11, bold=True, color="1A1A1A")
    notes_hdr.fill = hex_fill("F5F3F0")
    notes_hdr.border = Border(bottom=Side(style="medium", color="C4A484"))

    notes_items = [
        "The Status and Creator Tier columns have dropdown menus — click the cell to select a value.",
        "Conditional formatting is applied automatically to Status (column T) and Creator Tier (column U).",
        "Check and update this spreadsheet every day. No row should be left stale.",
        "Tier upgrades: Cold -> Warm after first content delivery. Warm -> Reliable after 2+ quality posts. Reliable -> VIP by management decision.",
    ]
    for i, note in enumerate(notes_items):
        r = 22 + i
        bullet_cell = leg.cell(row=r, column=2, value=f"  \u2022  ")
        bullet_cell.font = Font(name="Calibri", size=11, bold=True, color="C4A484")
        note_cell = leg.cell(row=r, column=3, value=note)
        note_cell.font = Font(name="Calibri", size=10, color="6B6B6B")
        note_cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
        leg.row_dimensions[r].height = 28

    # Save
    out_path = os.path.join(OUTPUT_DIR, "UGC Tracker.xlsx")
    wb.save(out_path)
    print(f"XLSX saved: {out_path}")
    return out_path


# ─────────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print("Generating UGC SOP DOCX...")
    build_docx()
    print("Generating UGC Tracker XLSX...")
    build_xlsx()
    print("Done.")
